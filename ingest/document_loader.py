"""Ekstrakcja tekstu z dokumentu załączonego przez użytkownika (PDF / DOCX / TXT).

Używane przez UI, gdy prawnik wgrywa pismo/umowę do analizy w kontekście przepisów.
"""
import io

import fitz  # pymupdf

MAX_CHARS = 60_000  # ostrożny limit, by nie wysyłać gigantycznych dokumentów do modelu
SUPPORTED = (".pdf", ".docx", ".txt")


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, bool]:
    """Zwraca (tekst, truncated). Rzuca ValueError dla nieobsługiwanego formatu."""
    name = filename.lower()
    if name.endswith(".pdf"):
        text = _from_pdf(file_bytes)
    elif name.endswith(".docx"):
        text = _from_docx(file_bytes)
    elif name.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Nieobsługiwany format pliku: {filename}. Dozwolone: {', '.join(SUPPORTED)}")

    text = text.strip()
    if len(text) > MAX_CHARS:
        return text[:MAX_CHARS], True
    return text, False


def _from_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = [page.get_text() for page in doc]
    doc.close()
    return "\n".join(pages)


def _from_docx(file_bytes: bytes) -> str:
    from docx import Document  # import lokalny — zależność opcjonalna
    document = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
